"""
Word Yardımcı Modülü
====================
Tüm Word belgelerinde kullanılan ortak fonksiyonlar.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def sayfa_ayarla(doc, sol=2.5, sag=2.0, ust=2.0, alt=2.0):
    for b in doc.sections:
        b.left_margin   = Cm(sol)
        b.right_margin  = Cm(sag)
        b.top_margin    = Cm(ust)
        b.bottom_margin = Cm(alt)


def hucre_boya(hucre, hex_renk: str):
    tc = hucre._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_renk)
    tcPr.append(shd)


def yeni_belge() -> Document:
    doc = Document()
    sayfa_ayarla(doc)
    return doc


def baslik_ekle(doc, metin: str, boyut=16, renk="1A3A5C",
                hizalama=WD_ALIGN_PARAGRAPH.CENTER, once=6, sonra=4):
    p = doc.add_paragraph()
    p.alignment = hizalama
    p.paragraph_format.space_before = Pt(once)
    p.paragraph_format.space_after  = Pt(sonra)
    run = p.add_run(metin)
    run.bold = True
    run.font.size = Pt(boyut)
    run.font.color.rgb = RGBColor.from_string(renk)
    run.font.name = "Arial"
    return p


def paragraf_ekle(doc, metin: str, boyut=10, kalin=False, renk="000000",
                  hizalama=WD_ALIGN_PARAGRAPH.LEFT, once=0, sonra=4):
    p = doc.add_paragraph()
    p.alignment = hizalama
    p.paragraph_format.space_before = Pt(once)
    p.paragraph_format.space_after  = Pt(sonra)
    run = p.add_run(metin)
    run.bold = kalin
    run.font.size = Pt(boyut)
    run.font.color.rgb = RGBColor.from_string(renk)
    run.font.name = "Arial"
    return p


def bolum_basligi(doc, metin: str, renk="1A3A5C"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"■  {metin}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(renk)
    run.font.name = "Arial"


def tablo_olustur(doc, basliklar: list, satirlar: list,
                  baslik_rengi="1A3A5C", satir_renkleri=("EBF3FB","FFFFFF")):
    """
    Genel amaçlı tablo üretir.
    basliklar: ["Sütun1", "Sütun2", ...]
    satirlar:  [["val1","val2",...], ...]
    """
    tablo = doc.add_table(rows=1 + len(satirlar), cols=len(basliklar))
    tablo.style = "Table Grid"

    # Başlık satırı
    for ci, b in enumerate(basliklar):
        h = tablo.rows[0].cells[ci]
        hucre_boya(h, baslik_rengi)
        run = h.paragraphs[0].add_run(b)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Arial"
        h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Veri satırları
    for ri, satir_veri in enumerate(satirlar):
        renk = satir_renkleri[ri % 2]
        for ci, deger in enumerate(satir_veri):
            h = tablo.rows[ri + 1].cells[ci]
            hucre_boya(h, renk)
            run = h.paragraphs[0].add_run(str(deger) if deger else "")
            run.font.size = Pt(9)
            run.font.name = "Arial"
            h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return tablo


def imza_tablosu(doc, etiketler: list, renk="1A3A5C"):
    """Altta imza bölümü ekler."""
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    tablo = doc.add_table(rows=2, cols=len(etiketler))
    tablo.style = "Table Grid"
    for ci, e in enumerate(etiketler):
        h = tablo.rows[0].cells[ci]
        hucre_boya(h, renk)
        run = h.paragraphs[0].add_run(e)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = "Arial"
        hucre_boya(tablo.rows[1].cells[ci], "FAFAFA")
        r2 = tablo.rows[1].cells[ci].paragraphs[0].add_run("\n\nİmza: ____________________")
        r2.font.size = Pt(9); r2.font.name = "Arial"
